import sys
import os
import traceback
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document

# Path to Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Tesseract language data directory
os.environ['TESSDATA_PREFIX'] = r'C:\Program Files\Tesseract-OCR\tessdata'

# Poppler path (used by pdf2image)
POPPLER_PATH = r'C:\tools\poppler-23.11.0\Library\bin'  # ← Replace with your actual Poppler path

def extract_text_from_pdf(pdf_path):
    """
    Convert each page of the PDF to an image and extract Arabic text using OCR.
    """
    try:
        images = convert_from_path(pdf_path, dpi=300, poppler_path=POPPLER_PATH)
        extracted_lines = []

        for idx, image in enumerate(images):
            print(f"Processing page {idx + 1}")
            text = pytesseract.image_to_string(image, config='--psm 6', lang='ara')
            extracted_lines.extend(text.splitlines())

        return extracted_lines

    except Exception as e:
        print("Error during PDF processing:", str(e))
        traceback.print_exc()
        return []

def save_to_word(data, output_file):
    """
    Save the extracted lines to a Word (.docx) file.
    """
    try:
        doc = Document()
        doc.add_heading('Extracted Arabic Text', level=1)

        for line in data:
            if line.strip():
                doc.add_paragraph(line)

        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        doc.save(output_file)
        print(f"Word file saved to: {output_file}")

    except Exception as e:
        print("Error during saving to Word:", str(e))
        traceback.print_exc()

if __name__ == "__main__":
    try:
        if len(sys.argv) < 2:
            print("Usage: python image_pro2.py <pdf_path>")
            sys.exit(1)

        pdf_path = sys.argv[1]

        if not os.path.exists(pdf_path):
            print(f"File not found: {pdf_path}")
            sys.exit(1)

        data = extract_text_from_pdf(pdf_path)

        if not data:
            print("No data extracted from the PDF.")
            sys.exit(1)

        output_file = os.path.join('storage', 'app', 'public', 'output.docx')
        save_to_word(data, output_file)

    except Exception as e:
        print("Error:", str(e))
        traceback.print_exc()
        sys.exit(1)
